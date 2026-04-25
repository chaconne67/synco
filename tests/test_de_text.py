import shutil

import pytest
from unittest.mock import patch
from data_extraction.services.text import extract_text


class TestExtractText:
    def test_docx_extraction(self, tmp_path):
        """python-docx로 .docx 파일에서 텍스트 추출."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("강솔찬")
        doc.add_paragraph("현대엠시트 회계팀장")
        filepath = tmp_path / "test.docx"
        doc.save(filepath)
        text = extract_text(str(filepath))
        assert "강솔찬" in text
        assert "현대엠시트" in text

    @patch("data_extraction.services.text._extract_doc_libreoffice")
    def test_doc_extraction_libreoffice_success(self, mock_libre, tmp_path):
        mock_libre.return_value = "강솔찬\n현대엠시트 회계팀장"
        filepath = tmp_path / "test.doc"
        filepath.write_bytes(b"fake doc content")
        text = extract_text(str(filepath))
        assert "강솔찬" in text
        mock_libre.assert_called_once()

    @patch("data_extraction.services.text._extract_doc_antiword")
    @patch("data_extraction.services.text._extract_doc_libreoffice")
    def test_doc_fallback_to_antiword(self, mock_libre, mock_antiword, tmp_path):
        mock_libre.side_effect = RuntimeError("libreoffice failed")
        mock_antiword.return_value = "강솔찬 이력서"
        filepath = tmp_path / "test.doc"
        filepath.write_bytes(b"fake doc content")
        text = extract_text(str(filepath))
        assert "강솔찬" in text
        mock_antiword.assert_called_once()

    @patch("data_extraction.services.text._extract_doc_antiword")
    @patch("data_extraction.services.text._extract_doc_libreoffice")
    def test_doc_fallback_to_antiword_when_libreoffice_returns_blankish(
        self,
        mock_libre,
        mock_antiword,
        tmp_path,
    ):
        mock_libre.return_value = "\ufeff\n\n\n"
        mock_antiword.return_value = "고영호 이력서"
        filepath = tmp_path / "test.doc"
        filepath.write_bytes(b"fake doc content")

        text = extract_text(str(filepath))

        assert text == "고영호 이력서"
        mock_antiword.assert_called_once()

    def test_pdf_extraction(self, tmp_path):
        """PyMuPDF로 .pdf 파일에서 텍스트 추출."""
        import fitz

        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        # Use ASCII text to avoid font/encoding issues in CI
        page.insert_text((72, 72), "Quality Manager Job Description 2026")
        doc.save(str(pdf_path))
        doc.close()

        text = extract_text(str(pdf_path))
        assert "Quality Manager" in text

    def test_unsupported_extension(self, tmp_path):
        filepath = tmp_path / "test.hwp"
        filepath.write_bytes(b"fake")
        with pytest.raises(ValueError, match="지원하지 않는"):
            extract_text(str(filepath))

    def test_empty_extraction(self, tmp_path):
        from docx import Document

        doc = Document()
        filepath = tmp_path / "empty.docx"
        doc.save(filepath)
        text = extract_text(str(filepath))
        assert text.strip() == ""


class TestBirthYearTextFilter:
    def test_extract_birth_year_from_dob(self):
        from data_extraction.services.text import extract_birth_year_from_text

        result = extract_birth_year_from_text("생년월일: 1985.03.12")

        assert result["birth_year"] == 1985
        assert result["source"] == "text_dob"

    def test_extract_birth_year_from_rrn(self):
        from data_extraction.services.text import extract_birth_year_from_text

        result = extract_birth_year_from_text("주민등록번호 850312-1******")

        assert result["birth_year"] == 1985
        assert result["source"] == "resident_registration_number"

    def test_two_digit_age_converts_to_cutoff_year(self):
        from datetime import date

        from data_extraction.services.text import normalize_birth_year_filter_value

        assert normalize_birth_year_filter_value(41, today=date(2026, 4, 25)) == 1985

    def test_passes_birth_year_filter_with_four_digit_cutoff(self):
        from data_extraction.services.text import passes_birth_year_filter

        result = passes_birth_year_filter(
            "생년월일: 1986.01.01",
            1985,
            enabled=True,
        )

        assert result.passed is True
        assert result.detected_year == 1986
        assert result.cutoff_year == 1985

    def test_birth_year_filter_blocks_older_candidate(self):
        from data_extraction.services.text import passes_birth_year_filter

        result = passes_birth_year_filter(
            "1984년생",
            1985,
            enabled=True,
        )

        assert result.passed is False
        assert result.detected_year == 1984


class TestExtractTextLibreoffice:
    @pytest.mark.skipif(
        shutil.which("libreoffice") is None,
        reason="LibreOffice not installed",
    )
    def test_extract_text_libreoffice_with_docx(self, tmp_path):
        from docx import Document

        from data_extraction.services.text import extract_text_libreoffice

        doc = Document()
        doc.add_paragraph("테스트 이력서 본문입니다")
        path = str(tmp_path / "test.docx")
        doc.save(path)
        result = extract_text_libreoffice(path)
        assert "테스트" in result
