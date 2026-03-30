import pytest
from unittest.mock import patch
from candidates.services.text_extraction import extract_text


class TestExtractText:
    def test_docx_extraction(self, tmp_path):
        """python-docx로 .docx 파일에서 텍스트 추출."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("강솔찬")
        doc.add_paragraph("현대엠시트 회계팀장")
        filepath = tmp_path / "test.docx"
        doc.save(filepath)
        text = extract_text(str(filepath))
        assert "강솔찬" in text
        assert "현대엠시트" in text

    @patch("candidates.services.text_extraction._extract_doc_antiword")
    def test_doc_extraction_antiword_success(self, mock_antiword, tmp_path):
        mock_antiword.return_value = "강솔찬\n현대엠시트 회계팀장"
        filepath = tmp_path / "test.doc"
        filepath.write_bytes(b"fake doc content")
        text = extract_text(str(filepath))
        assert "강솔찬" in text
        mock_antiword.assert_called_once()

    @patch("candidates.services.text_extraction._extract_doc_libreoffice")
    @patch("candidates.services.text_extraction._extract_doc_antiword")
    def test_doc_fallback_to_libreoffice(self, mock_antiword, mock_libre, tmp_path):
        mock_antiword.side_effect = RuntimeError("antiword failed")
        mock_libre.return_value = "강솔찬 이력서"
        filepath = tmp_path / "test.doc"
        filepath.write_bytes(b"fake doc content")
        text = extract_text(str(filepath))
        assert "강솔찬" in text
        mock_libre.assert_called_once()

    def test_unsupported_extension(self, tmp_path):
        filepath = tmp_path / "test.hwp"
        filepath.write_bytes(b"fake")
        with pytest.raises(ValueError, match="지원하지 않는"):
            extract_text(str(filepath))

    def test_empty_extraction(self, tmp_path):
        from docx import Document

        doc = Document()
        filepath = tmp_path / "empty.docx"
        doc.save(filepath)
        text = extract_text(str(filepath))
        assert text == ""
