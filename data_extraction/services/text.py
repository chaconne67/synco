"""Extract text from .doc and .docx resume files."""

import os
import re
import subprocess
import tempfile
from dataclasses import dataclass
from datetime import date

from docx import Document


def extract_text(file_path: str) -> str:
    """Extract text from a resume file based on its extension.

    Supports .docx (python-docx), .doc (antiword / LibreOffice fallback),
    and .pdf (PyMuPDF).
    Raises ValueError for unsupported file formats.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".doc":
        return _extract_doc(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}")


def _extract_docx(file_path: str) -> str:
    """Extract text from a .docx file.

    Tries python-docx first (paragraphs, tables, textboxes).
    If the result is too short, tries LibreOffice and picks the longer result.
    This handles cases where python-docx silently misses table data.
    """
    docx_text = ""
    try:
        doc = Document(file_path)
        parts: list[str] = []

        # 1) Body paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # 2) Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        # 3) Textboxes (VML w:txbxContent) — often contain headers with
        #    company names, positions, and date ranges in Korean resumes
        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        seen: set[str] = set()
        for txbx in doc.element.body.findall(f".//{{{ns_w}}}txbxContent"):
            t_elems = txbx.findall(f".//{{{ns_w}}}t")
            text = " ".join((t.text or "") for t in t_elems).strip()
            if text and text not in seen:
                seen.add(text)
                parts.append(text)

        docx_text = "\n".join(parts)
    except Exception:
        pass

    # Try LibreOffice and pick the result with richer content.
    # Length alone is not a reliable indicator — a longer result may still
    # miss table data while a shorter result captures it.
    try:
        lo_text = _extract_doc_libreoffice(file_path)
    except Exception:
        lo_text = ""

    if not _strip_bom(lo_text):
        return docx_text or ""
    if not _strip_bom(docx_text):
        return lo_text

    # Pick the result with more structural resume signals
    docx_score = _content_richness_score(docx_text)
    lo_score = _content_richness_score(lo_text)
    return lo_text if lo_score > docx_score else docx_text


def _extract_doc(file_path: str) -> str:
    """Extract text from a .doc file. Tries both LibreOffice and antiword,
    picks the result with richer content."""
    lo_text = ""
    aw_text = ""

    try:
        lo_text = _extract_doc_libreoffice(file_path)
    except Exception:
        pass

    try:
        aw_text = _extract_doc_antiword(file_path)
    except Exception:
        pass

    if not _strip_bom(lo_text) and not _strip_bom(aw_text):
        return ""
    if not _strip_bom(lo_text):
        return aw_text
    if not _strip_bom(aw_text):
        return lo_text

    lo_score = _content_richness_score(lo_text)
    aw_score = _content_richness_score(aw_text)
    return lo_text if lo_score >= aw_score else aw_text


def _extract_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: uv add pymupdf")

    doc = fitz.open(file_path)
    parts: list[str] = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            parts.append(text)
    doc.close()
    return "\n".join(parts)


def _extract_doc_antiword(file_path: str) -> str:
    """Extract text from a .doc file using antiword."""
    result = subprocess.run(
        ["antiword", file_path],
        capture_output=True,
        timeout=30,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"antiword failed with return code {result.returncode}: "
            f"{result.stderr.decode('utf-8', errors='replace')}"
        )

    # Try decoding with multiple encodings
    for encoding in ("utf-8", "euc-kr", "cp949", "latin-1"):
        try:
            return result.stdout.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue

    # latin-1 should never fail, but just in case
    return result.stdout.decode("latin-1")


def _extract_doc_libreoffice(file_path: str) -> str:
    """Extract text from a .doc file using LibreOffice headless conversion."""
    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                tmpdir,
                file_path,
            ],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice failed: {result.stderr.decode('utf-8', errors='replace')}"
            )

        # Find the converted .txt file
        basename = os.path.splitext(os.path.basename(file_path))[0]
        txt_path = os.path.join(tmpdir, f"{basename}.txt")

        if not os.path.exists(txt_path):
            raise RuntimeError(f"LibreOffice conversion failed: {txt_path} not found")

        # Try decoding with multiple encodings
        with open(txt_path, "rb") as f:
            raw = f.read()
        for encoding in ("utf-8", "euc-kr", "cp949", "latin-1"):
            try:
                return raw.decode(encoding)
            except (UnicodeDecodeError, ValueError):
                continue

        # latin-1 should never fail, but just in case
        return raw.decode("latin-1")


# Lines that are nothing but a personal-info form label without a value
# (e.g. "이름:", "Email:", "주소"). LLM 입력에서 정보 가치 0.
#
# IMPORTANT: do NOT include section headers like "학력/Education", "경력/Career",
# "자격증/Certification", "어학/Language" here — those signal the start of a
# structured section to both the LLM and the self-intro region detector. If we
# strip them, _compress_self_intro_region loses its exit cue and ends up
# eating real Education/Career rows.
_FORM_LABEL_ONLY = re.compile(
    r"^\s*(?:이름|성명|name|성별|gender|나이|age|연락처|tel|phone|email|"
    r"이메일|주소|address|특기|취미|병역|military|"
    r"본적|현주소|등록기준지)"
    r"\s*[:：]?\s*$",
    re.IGNORECASE,
)

# Korean resume signature line at end of document
# (e.g. "2024년 3월 22일 / 성명: 홍길동 (인)")
_SIGNATURE_DATE_LINE = re.compile(
    r"^\s*\d{4}\s*[년.\-/]\s*\d{1,2}\s*[월.\-/]\s*\d{1,2}\s*일?\s*$"
)
_SIGNATURE_NAME_LINE = re.compile(
    r"\b(?:서명|성명|날인|signature)\s*[:：]?\s*[가-힣A-Za-z\s]{2,15}\s*[\(（]\s*인\s*[\)）]",
    re.IGNORECASE,
)

# Self-introduction / cover-letter section headers — content after these
# headers is mostly free prose with low signal-density for structured fields.
_SELF_INTRO_HEADER = re.compile(
    r"자기\s*소개(?:서)?|지원\s*동기|성장\s*과정|입사\s*포부|"
    r"성격(?:의)?\s*장단점|장단점|"
    r"personal\s*statement|cover\s*letter|career\s*objective|"
    r"about\s+me|professional\s+summary",
    re.IGNORECASE,
)

# Headers that mark the END of a self-intro region. When we hit one of these,
# we resume verbatim copying — protects Education/Career/Certificates that
# follow a free-prose About-Me block.
_STRUCTURED_SECTION_HEADER = re.compile(
    r"^\s*(?:"
    r"학력(?:\s*사항)?|education(?:al\s+background)?|"
    r"경력(?:\s*사항)?|경력\s*기술서|career(?:\s+history|\s+summary)?|"
    r"experience|work\s*history|employment(?:\s+history)?|"
    r"professional\s+experience|employment(?:\s+record)?|"
    r"자격(?:\s*사항|\s*증)?|certification|license|qualification|certificates?|"
    r"기술(?:\s*스택)?|skills?|expertise|technical\s+skills?|core\s+competenc(?:y|ies)|"
    r"어학(?:\s*능력)?|language(?:\s+skills?)?|"
    r"수상(?:\s*경력)?|awards?|honors?|"
    r"출판|publications?|논문|"
    r"프로젝트|projects?|"
    r"활동|activities|extracurricular|"
    r"참고|references?|"
    r"교육(?:\s*이수)?|training|courses?"
    r")\s*[:：]?\s*$",
    re.IGNORECASE,
)

# Lines worth preserving inside the self-intro region: company names,
# year markers, education institutions. Anything matched here keeps the
# integrity-comparison signal alive even after compression.
_COMPANY_OR_YEAR_LINE = re.compile(
    # 4자리 연도
    r"(?:19|20)\d{2}|"
    # 한국식 회사 표기
    r"㈜|\(주\)|주식회사|"
    # 영문 회사 접미사
    r"\b(?:co\.|corp\.|inc\.|llc|gmbh|ltd\.|company|group)\b|"
    # 한국어 회사·기관 접미사가 붙은 단어
    r"[가-힣]{2,}(?:전자|화학|중공업|건설|상사|컴퍼니|코리아|은행|증권|카드|보험|"
    r"제약|병원|연구소|대학교|대학원|대학|학교|회사|호텔|에너지|시스템|텔레콤|"
    r"커뮤니케이션|엔터테인먼트|네트워크|솔루션|테크|레저|미디어|모터스)",
    re.IGNORECASE,
)


def _drop_form_label_lines(lines: list[str]) -> list[str]:
    return [ln for ln in lines if not _FORM_LABEL_ONLY.match(ln)]


def _drop_trailing_signature(lines: list[str]) -> list[str]:
    """Drop signature/date lines that appear in the last 5 lines of the document.

    Restricting to the tail avoids removing legitimate dates from the body
    (career start dates, certification dates, etc.).
    """
    if len(lines) < 2:
        return lines
    out = list(lines)
    tail_window = 5
    for _ in range(tail_window):
        if not out:
            break
        last = out[-1]
        if _SIGNATURE_DATE_LINE.match(last) or _SIGNATURE_NAME_LINE.search(last):
            out.pop()
            continue
        break
    return out


def _merge_fragmented_lines(lines: list[str]) -> list[str]:
    """Merge consecutive very-short lines (table-cell fragmentation).

    Heuristic: if a line is < 5 non-whitespace chars and the next line is also
    short (< 30 chars), join them with a single space. This recovers semantic
    units that .doc/.docx extraction split into per-cell rows.
    """
    if len(lines) < 2:
        return lines
    out: list[str] = []
    i = 0
    while i < len(lines):
        cur = lines[i]
        cur_short = len(cur.strip()) < 5
        if cur_short and i + 1 < len(lines):
            nxt = lines[i + 1]
            if len(nxt.strip()) < 30:
                out.append(f"{cur} {nxt}".strip())
                i += 2
                continue
        out.append(cur)
        i += 1
    return out


def _compress_self_intro_region(lines: list[str]) -> list[str]:
    """Compress free-prose self-intro regions while preserving structured sections.

    State machine over lines:
      - default: copy verbatim. Switch to compressing when a self-intro header
        appears.
      - compressing: drop lines unless they mention a company/year (keeps
        integrity signals). Exit back to verbatim when a structured-section
        header appears (Education / Career / Skills / etc.) — this prevents
        accidentally swallowing the structured part of the resume that comes
        after a short About-Me block.
    """
    out: list[str] = []
    compressing = False
    skipped = 0

    def flush_skipped():
        nonlocal skipped
        if skipped > 0:
            out.append(f"[자기소개 산문 {skipped}줄 생략]")
            skipped = 0

    for ln in lines:
        if compressing:
            if _STRUCTURED_SECTION_HEADER.match(ln):
                flush_skipped()
                compressing = False
                out.append(ln)
                continue
            if _COMPANY_OR_YEAR_LINE.search(ln):
                flush_skipped()
                out.append(ln)
            else:
                skipped += 1
            continue

        # not compressing
        if _SELF_INTRO_HEADER.search(ln):
            out.append(ln)
            compressing = True
            continue
        out.append(ln)

    flush_skipped()
    return out


def preprocess_resume_text(text: str) -> str:
    """Clean and deduplicate resume text to reduce LLM token usage.

    Pipeline:
      1) sanitize encoding/control chars + Word field codes
      2) remove blank lines, compress whitespace
      3) remove exact-duplicate lines
      4) remove near-duplicate lines (70% word overlap)
      5) drop generic noise lines (PC skills only, no company/date)
      6) drop form-label-only lines
      7) drop trailing signature/date lines
      8) merge fragmented short lines (table cell recovery)
      9) compress self-intro region — keep only company/year lines
    """
    from data_extraction.services.extraction.sanitizers import sanitize_input_text

    text = sanitize_input_text(text)
    lines = text.split("\n")

    # 1) Remove blank lines, compress whitespace
    lines = [re.sub(r"\s{2,}", " ", line).strip() for line in lines if line.strip()]

    # 2) Remove exact duplicate lines (preserve order)
    seen: set[str] = set()
    unique: list[str] = []
    for line in lines:
        normalized = line.strip().lower()
        if normalized and normalized not in seen:
            seen.add(normalized)
            unique.append(line)

    # 3) Remove near-duplicate lines (70%+ word overlap with recent lines)
    final: list[str] = []
    for line in unique:
        words = set(line.lower().split())
        if len(words) < 3:
            final.append(line)
            continue
        is_dup = False
        for existing in final[-10:]:
            existing_words = set(existing.lower().split())
            if existing_words:
                overlap = len(words & existing_words) / max(
                    len(words), len(existing_words)
                )
                if overlap > 0.7:
                    is_dup = True
                    break
        if not is_dup:
            final.append(line)

    # 4) Remove noise patterns (basic PC skills, etc.)
    noise = [
        "워드/엑셀",
        "ms-office",
        "ms office",
        "powerpoint",
        "computer :",
        "computer:",
    ]
    date_pattern = re.compile(r"\d{4}")
    company_suffixes = ("㈜", "주식회사", "(주)", "co.", "corp", "inc")
    cleaned: list[str] = []
    for line in final:
        lower = line.lower()
        if any(n in lower for n in noise):
            if date_pattern.search(line) or any(s in lower for s in company_suffixes):
                cleaned.append(line)
                continue
            if len(line.strip()) < 40:
                continue
        cleaned.append(line)
    final = cleaned

    # 5) New conservative passes
    final = _drop_form_label_lines(final)
    final = _drop_trailing_signature(final)
    final = _merge_fragmented_lines(final)
    final = _compress_self_intro_region(final)

    return "\n".join(final)


def _content_richness_score(text: str) -> int:
    """Score text by how many resume-structural signals it contains.

    Checks for presence of career/education keywords, contact info patterns,
    and date patterns — not just length.
    """
    score = 0
    t = text.lower()

    # Career signals
    for kw in ["경력", "재직", "회사명", "근무기간", "experience", "career"]:
        if kw in t:
            score += 1

    # Education signals
    for kw in ["학력", "대학", "졸업", "university", "education"]:
        if kw in t:
            score += 1

    # Contact signals
    if re.search(r"010[-.\s]?\d{4}[-.\s]?\d{4}", text):
        score += 2
    if re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text):
        score += 2

    # Date patterns (YYYY.MM or YYYY-MM)
    date_count = len(re.findall(r"\d{4}[.-/]\d{1,2}", text))
    score += min(date_count, 5)  # cap at 5

    # Company name patterns (㈜, (주))
    company_count = len(re.findall(r"㈜|\(주\)|주식회사", text))
    score += min(company_count, 3)

    return score


def _strip_bom(text: str) -> str:
    """Strip BOM and whitespace."""
    return text.replace("\ufeff", "").strip()


def _has_substantive_text(text: str) -> bool:
    if not text:
        return False
    return bool(re.search(r"[0-9A-Za-z\uac00-\ud7a3]", text))


def classify_text_quality(text: str) -> str:
    """Classify extracted text quality before LLM extraction.

    Returns: 'ok', 'too_short', 'garbled', 'empty'
    """
    if not text or not text.strip():
        return "empty"

    stripped = text.strip()

    # Check ratio of meaningful characters (Korean, Latin, digits)
    alnum_chars = sum(1 for c in stripped if c.isalnum() or "\uac00" <= c <= "\ud7a3")
    if len(stripped) > 0 and alnum_chars / len(stripped) < 0.3:
        return "garbled"

    # Reject truly empty content only — actual quality is judged post-extraction
    if len(stripped) < 50:
        return "too_short"

    return "ok"


@dataclass(frozen=True)
class BirthYearFilterResult:
    passed: bool
    active: bool
    cutoff_year: int | None = None
    detected_year: int | None = None
    source: str = ""
    evidence: str = ""
    reason: str = ""


def normalize_birth_year_filter_value(value: int, *, today: date | None = None) -> int:
    """Normalize a 2-digit age or 4-digit birth year into a cutoff birth year."""
    today = today or date.today()
    if 10 <= value <= 99:
        return today.year - value
    if 1900 <= value <= today.year:
        return value
    raise ValueError("birth-year filter must be a 2-digit age or 4-digit birth year")


def extract_birth_year_from_text(text: str, *, today: date | None = None) -> dict | None:
    """Extract a likely birth year from resume text using contextual patterns."""
    today = today or date.today()
    if not text:
        return None

    patterns = [
        (
            "text_dob",
            re.compile(
                r"(?:생년월일|출생(?:년도|연도)?|출생일|birthday|birth|d\.?o\.?b\.?|"
                r"date\s+of\s+birth|born)"
                r"[\s:：\-]*"
                r"(?P<year>(?:19|20)\d{2})[.\-/년\s]*(?:\d{1,2})?",
                re.IGNORECASE,
            ),
        ),
        (
            # "1981년생" 또는 "1981년 1월 9일생" — 한국어 표준 표기
            "text_birth_year",
            re.compile(
                r"(?P<year>(?:19|20)\d{2})\s*년"
                r"(?:\s*\d{1,2}\s*월(?:\s*\d{1,2}\s*일)?)?"
                r"\s*생"
            ),
        ),
        (
            # "1981.01.09생" / "1981-01-09 생" / "1981/01/09생"
            "text_date_birth",
            re.compile(
                r"(?P<year>(?:19|20)\d{2})[.\-/]\d{1,2}[.\-/]\d{1,2}\s*생"
            ),
        ),
        (
            # "D.O.B. 01.09.1981" / "DOB: 12/15/1985" — 영문 DMY 형식
            "text_dob_dmy",
            re.compile(
                r"(?:d\.?o\.?b\.?|date\s+of\s+birth|birthday|born)"
                r"[\s:：\-]*"
                r"\d{1,2}[.\-/]\d{1,2}[.\-/](?P<year>(?:19|20)\d{2})",
                re.IGNORECASE,
            ),
        ),
        (
            "text_short_birth_year",
            re.compile(r"(?P<yy>\d{2})\s*년\s*생"),
        ),
    ]
    for source, pattern in patterns:
        match = pattern.search(text)
        if not match:
            continue
        year = _normalize_birth_match(match, today=today)
        if year:
            return {
                "birth_year": year,
                "source": source,
                "evidence": match.group(0).strip()[:120],
            }

    rrn = re.search(r"(?<!\d)(?P<yy>\d{2})(?P<mm>0[1-9]|1[0-2])(?P<dd>[0-3]\d)[-\s]?(?P<gender>[1-4])", text)
    if rrn:
        yy = int(rrn.group("yy"))
        gender = rrn.group("gender")
        year = (1900 + yy) if gender in {"1", "2"} else (2000 + yy)
        if 1900 <= year <= today.year:
            return {
                "birth_year": year,
                "source": "resident_registration_number",
                "evidence": rrn.group(0).strip()[:120],
            }

    age_match = re.search(
        r"(?:나이|연령|age)[\s:：\-]*(?:만\s*)?(?P<age>\d{2})\s*(?:세|years?\s+old)?",
        text,
        re.IGNORECASE,
    )
    if age_match:
        age = int(age_match.group("age"))
        if 10 <= age <= 99:
            return {
                "birth_year": today.year - age,
                "source": "age",
                "evidence": age_match.group(0).strip()[:120],
            }

    return None


def passes_birth_year_filter(
    text: str,
    filter_value: int | None,
    *,
    enabled: bool = False,
    today: date | None = None,
    file_name: str | None = None,
) -> BirthYearFilterResult:
    """Return whether text passes the pre-LLM birth-year cutoff filter.

    Lookup order:
    1. Resume body (regex patterns over preprocessed text).
    2. Filename (only when text lookup misses) — Korean resume convention
       often encodes year in filename like "이름.85.회사.학교.docx".
    3. If both fail, pass conservatively. Dropping a resume on missing data
       has higher false-negative cost than processing one extra resume.
    """
    if not enabled or filter_value is None:
        return BirthYearFilterResult(passed=True, active=False)

    cutoff = normalize_birth_year_filter_value(filter_value, today=today)
    extracted = extract_birth_year_from_text(text, today=today)

    if not extracted and file_name:
        from data_extraction.services.filename import parse_filename

        parsed = parse_filename(file_name)
        if parsed.get("birth_year"):
            extracted = {
                "birth_year": parsed["birth_year"],
                "source": "filename",
                "evidence": file_name[:120],
            }

    if not extracted:
        return BirthYearFilterResult(
            passed=True,
            active=True,
            cutoff_year=cutoff,
            source="not_detected",
            reason="birth year not found; passed conservatively",
        )

    detected = extracted["birth_year"]
    passed = detected >= cutoff
    return BirthYearFilterResult(
        passed=passed,
        active=True,
        cutoff_year=cutoff,
        detected_year=detected,
        source=extracted["source"],
        evidence=extracted["evidence"],
        reason=(
            ""
            if passed
            else f"detected birth year {detected} is before cutoff {cutoff}"
        ),
    )


def _normalize_birth_match(match: re.Match, *, today: date) -> int | None:
    if match.groupdict().get("year"):
        year = int(match.group("year"))
    else:
        yy = int(match.group("yy"))
        year = 1900 + yy if yy >= 50 else 2000 + yy
    if 1900 <= year <= today.year:
        return year
    return None


def extract_text_libreoffice(file_path: str) -> str:
    """Public wrapper for LibreOffice-based text extraction.

    Used by retry pipeline when python-docx extraction misses content.
    """
    return _extract_doc_libreoffice(file_path)
