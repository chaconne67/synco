"""Extract text from .doc and .docx resume files."""

import os
import subprocess
import tempfile

from docx import Document


def extract_text(file_path: str) -> str:
    """Extract text from a resume file based on its extension.

    Supports .docx (python-docx) and .doc (antiword / LibreOffice fallback).
    Raises ValueError for unsupported file formats.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".doc":
        return _extract_doc(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식: {ext}")


def _extract_docx(file_path: str) -> str:
    """Extract text from a .docx file using python-docx.

    Extracts from paragraphs, tables, AND textboxes (VML/WPS).
    Falls back to LibreOffice if the file is corrupted or not a valid docx.
    """
    try:
        doc = Document(file_path)
        parts: list[str] = []

        # 1) Body paragraphs
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)

        # 2) Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)

        # 3) Textboxes (VML w:txbxContent) — often contain headers with
        #    company names, positions, and date ranges in Korean resumes
        ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        seen: set[str] = set()
        for txbx in doc.element.body.findall(f".//{{{ns_w}}}txbxContent"):
            t_elems = txbx.findall(f".//{{{ns_w}}}t")
            text = " ".join((t.text or "") for t in t_elems).strip()
            # Deduplicate (VML textboxes often appear twice — mc:Choice + mc:Fallback)
            if text and text not in seen:
                seen.add(text)
                parts.append(text)

        return "\n".join(parts)
    except Exception:
        # Corrupted/invalid docx — try LibreOffice as fallback
        return _extract_doc_libreoffice(file_path)


def _extract_doc(file_path: str) -> str:
    """Extract text from a .doc file. Tries antiword first, falls back to LibreOffice."""
    try:
        return _extract_doc_antiword(file_path)
    except Exception:
        return _extract_doc_libreoffice(file_path)


def _extract_doc_antiword(file_path: str) -> str:
    """Extract text from a .doc file using antiword."""
    result = subprocess.run(
        ["antiword", file_path],
        capture_output=True,
        timeout=30,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"antiword failed with return code {result.returncode}: "
            f"{result.stderr.decode('utf-8', errors='replace')}"
        )

    # Try decoding with multiple encodings
    for encoding in ("utf-8", "euc-kr", "cp949", "latin-1"):
        try:
            return result.stdout.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue

    # latin-1 should never fail, but just in case
    return result.stdout.decode("latin-1")


def _extract_doc_libreoffice(file_path: str) -> str:
    """Extract text from a .doc file using LibreOffice headless conversion."""
    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "txt:Text",
                "--outdir",
                tmpdir,
                file_path,
            ],
            capture_output=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice failed: {result.stderr.decode('utf-8', errors='replace')}"
            )

        # Find the converted .txt file
        basename = os.path.splitext(os.path.basename(file_path))[0]
        txt_path = os.path.join(tmpdir, f"{basename}.txt")

        if not os.path.exists(txt_path):
            raise RuntimeError(f"LibreOffice conversion failed: {txt_path} not found")

        # Try decoding with multiple encodings
        with open(txt_path, "rb") as f:
            raw = f.read()
        for encoding in ("utf-8", "euc-kr", "cp949", "latin-1"):
            try:
                return raw.decode(encoding)
            except (UnicodeDecodeError, ValueError):
                continue

        # latin-1 should never fail, but just in case
        return raw.decode("latin-1")


def extract_text_libreoffice(file_path: str) -> str:
    """Public wrapper for LibreOffice-based text extraction.

    Used by retry pipeline when python-docx extraction misses content.
    """
    return _extract_doc_libreoffice(file_path)
