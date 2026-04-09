"""Word 변환 + 마스킹 처리 (python-docx)."""

import copy
import io
import logging

from django.core.files.base import ContentFile
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from projects.models import DEFAULT_MASKING_CONFIG

logger = logging.getLogger(__name__)

# 마스킹 대상 필드 매핑
MASKING_FIELDS = {
    "salary": [
        "current_salary",
        "desired_salary",
        "salary_detail",
        "salary_expectation",
    ],
    "birth_detail": ["birth_year"],
    "contact": ["email", "phone", "address"],
    "current_company": ["current_company"],
}


def _apply_masking(data: dict, masking_config: dict) -> dict:
    """마스킹 설정에 따라 데이터에서 민감 필드를 제거."""
    masked = copy.deepcopy(data)

    for mask_key, should_mask in masking_config.items():
        if not should_mask:
            continue
        fields = MASKING_FIELDS.get(mask_key, [])
        for field in fields:
            # top-level
            if field in masked:
                masked[field] = "[마스킹]"
            # personal_info nested
            if "personal_info" in masked and field in masked["personal_info"]:
                masked["personal_info"][field] = "[마스킹]"
            # additional nested
            if "additional" in masked and field in masked.get("additional", {}):
                masked["additional"][field] = "[마스킹]"

    return masked


def _add_section(doc: Document, title: str, content) -> None:
    """Word 문서에 섹션 추가."""
    doc.add_heading(title, level=2)

    if isinstance(content, str):
        doc.add_paragraph(content)
    elif isinstance(content, list):
        for item in content:
            if isinstance(item, dict):
                for k, v in item.items():
                    if v and v != "[마스킹]":
                        doc.add_paragraph(f"{k}: {v}", style="List Bullet")
            else:
                doc.add_paragraph(str(item), style="List Bullet")
    elif isinstance(content, dict):
        for k, v in content.items():
            if v and v != "[마스킹]":
                doc.add_paragraph(f"{k}: {v}")


def _build_document(data: dict) -> Document:
    """final_content_json에서 Word 문서 생성."""
    doc = Document()

    # 스타일 설정
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    # 제목
    personal = data.get("personal_info", {})
    name = personal.get("name", "")
    name_en = personal.get("name_en", "")
    title = f"추천 서류 — {name}"
    if name_en and name_en != "[마스킹]":
        title += f" ({name_en})"
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 인적사항
    if personal:
        _add_section(doc, "인적사항", personal)

    # 요약
    if data.get("summary"):
        _add_section(doc, "전문 요약", data["summary"])

    # 핵심 역량
    if data.get("core_competencies"):
        _add_section(doc, "핵심 역량", data["core_competencies"])

    # 경력
    if data.get("careers"):
        doc.add_heading("경력사항", level=2)
        for career in data["careers"]:
            company = career.get("company", "")
            position = career.get("position", "")
            period = career.get("period", "")
            doc.add_heading(f"{company} — {position} ({period})", level=3)
            if career.get("company_intro"):
                doc.add_paragraph(career["company_intro"])
            if career.get("responsibilities"):
                for resp in career["responsibilities"]:
                    doc.add_paragraph(resp, style="List Bullet")

    # 학력
    if data.get("educations"):
        _add_section(doc, "학력", data["educations"])

    # 자격증
    if data.get("certifications"):
        _add_section(doc, "자격증/면허", data["certifications"])

    # 어학
    if data.get("language_skills"):
        _add_section(doc, "어학능력", data["language_skills"])

    # 기술
    if data.get("skills"):
        _add_section(doc, "보유 기술", data["skills"])

    # 병역
    if data.get("military"):
        military = data["military"]
        if any(v for v in military.values()):
            _add_section(doc, "병역", military)

    # 기타
    if data.get("additional"):
        additional = data["additional"]
        for key, label in [
            ("awards", "수상경력"),
            ("patents", "특허"),
            ("overseas", "해외경험"),
            ("training", "교육이수"),
        ]:
            if additional.get(key):
                _add_section(doc, label, additional[key])
        if additional.get("self_introduction"):
            _add_section(doc, "자기소개", additional["self_introduction"])

    return doc


def convert_to_word(draft) -> None:
    """Draft의 final_content_json을 Word 파일로 변환."""
    data = draft.final_content_json
    if not data:
        raise RuntimeError("최종 정리 데이터가 없습니다.")

    # 마스킹 적용
    masked_data = _apply_masking(data, draft.masking_config or DEFAULT_MASKING_CONFIG)

    # Word 문서 생성
    doc = _build_document(masked_data)

    # 파일로 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    candidate_name = draft.submission.candidate.name
    filename = f"추천서류_{candidate_name}.docx"
    draft.output_file.save(filename, ContentFile(buffer.read()), save=True)
